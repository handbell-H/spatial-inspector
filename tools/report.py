"""검수 결과 Word 보고서 생성"""

from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 스타일 헬퍼 ──────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _header_row(table, headers: list, bg="2E5A9C"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        _set_cell_bg(cell, bg)


def _add_row(table, values: list, center_cols: list = None):
    row = table.add_row()
    center_cols = center_cols or []
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v) if v is not None else "-"
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in center_cols else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(9)


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _para(doc, text, bold=False, size=10):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.bold = bold
        run.font.size = Pt(size)
    return p


# ── 각 섹션 ──────────────────────────────────────────────────

def _section_overview(doc, prev_folder, curr_folder, emd_shp):
    _heading(doc, "1. 검수 개요", level=1)
    today = date.today().strftime("%Y년 %m월 %d일")
    items = [
        ("검수 일자", today),
        ("전년도 자료 경로", prev_folder or "-"),
        ("금년도 자료 경로", curr_folder),
        ("읍면동 경계 파일", emd_shp or "-"),
        ("검수 항목", "① 시계열 수량 비교  ② 중복 레코드 탐지  ③ 지오코딩 정확도"),
        ("좌표계", "EPSG:5179 (Korea 2000 / Unified CS)"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _header_row(table, ["항목", "내용"])
    for k, v in items:
        _add_row(table, [k, v], center_cols=[0])
    doc.add_paragraph()


def _section_count(doc, count_data: dict):
    _heading(doc, "2. 시계열 시설별 수량 비교", level=1)
    s = count_data["summary"]
    _para(doc,
          f"전년도 시설 종류: {s['전년_시설종류수']}종  |  "
          f"금년도 시설 종류: {s['금년_시설종류수']}종  |  "
          f"전년 총수량: {s['전년_총수량']:,}건  →  "
          f"금년 총수량: {s['금년_총수량']:,}건  (총증감: {s['총증감']:+,}건)",
          size=9)

    if s["신규시설"]:
        _para(doc, f"▶ 신규 시설 종류: {', '.join(s['신규시설'])}", size=9)
    if s["추후수령후분석"]:
        _para(doc, f"▶ 추후 수령 후 분석 대상: {', '.join(s['추후수령후분석'])}", size=9)

    headers = ["시설명", "전년 수량", "금년 수량", "증감", "상태"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    for row in count_data["rows"]:
        diff_str = (f"{row['증감']:+,}" if isinstance(row["증감"], int) else "-")
        _add_row(table,
                 [row["시설명"],
                  f"{row['전년수량']:,}" if isinstance(row["전년수량"], int) else "-",
                  f"{row['금년수량']:,}" if isinstance(row["금년수량"], int) else "-",
                  diff_str,
                  row["상태"]],
                 center_cols=[1, 2, 3, 4])
    doc.add_paragraph()


def _section_duplicate(doc, dup_data: dict):
    _heading(doc, "3. 중복 레코드 현황", level=1)
    _para(doc,
          f"검사 파일 수: {dup_data['검사파일수']}개  |  총 중복 레코드 수: {dup_data['총중복수']:,}건",
          size=9)
    _para(doc, "※ 중복 기준: fac_nm + fac_add + x_coord + y_coord 4개 컬럼 모두 동일한 레코드", size=8)

    # 시설별 요약 표
    headers = ["시설명", "전체수량", "중복수", "제거 후 수량"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    for r in dup_data["시설별결과"]:
        if "오류" in r:
            _add_row(table, [r["시설명"], "오류", "-", "-"], center_cols=[1, 2, 3])
        else:
            _add_row(table,
                     [r["시설명"],
                      f"{r['전체수량']:,}",
                      f"{r['중복수']:,}",
                      f"{r['제거후수량']:,}"],
                     center_cols=[1, 2, 3])

    # 중복 레코드 상세
    has_detail = any(r.get("중복레코드") for r in dup_data["시설별결과"])
    if has_detail:
        doc.add_paragraph()
        _heading(doc, "3-1. 중복 레코드 상세 (시설별 최대 20건)", level=2)
        for r in dup_data["시설별결과"]:
            if not r.get("중복레코드"):
                continue
            _para(doc, f"▶ {r['시설명']} ({r['중복수']}건)", bold=True, size=9)
            det_headers = ["시설명(fac_nm)", "주소(fac_add)", "X좌표", "Y좌표"]
            det_table = doc.add_table(rows=1, cols=len(det_headers))
            det_table.style = "Table Grid"
            _header_row(det_table, det_headers, bg="5B7FC0")
            for rec in r["중복레코드"]:
                _add_row(det_table,
                         [rec.get("fac_nm", "-"),
                          rec.get("fac_add", "-"),
                          rec.get("x_coord", "-"),
                          rec.get("y_coord", "-")],
                         center_cols=[2, 3])
    doc.add_paragraph()


def _section_geocode(doc, geo_data: dict):
    _heading(doc, "4. 지오코딩 정확도", level=1)
    _para(doc,
          f"전체 수량: {geo_data['전체수량']:,}건  |  "
          f"전체 일치: {geo_data['전체일치']:,}건  |  "
          f"전체 일치율: {geo_data['전체일치율(%)']:.1f}%",
          size=9)
    _para(doc,
          f"※ 읍면동 경계 컬럼: {geo_data['사용된_읍면동컬럼']}  |  "
          "검사 방법: point 좌표를 읍면동 경계 SHP와 공간조인 후 fac_add의 읍면동명과 비교",
          size=8)

    headers = ["시설명", "전체수량", "일치", "불일치", "조인실패", "일치율(%)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _header_row(table, headers)

    for r in geo_data["시설별결과"]:
        if "오류" in r:
            _add_row(table, [r["시설명"], "오류", "-", "-", "-", "-"], center_cols=list(range(6)))
        else:
            _add_row(table,
                     [r["시설명"],
                      f"{r['전체수량']:,}",
                      f"{r['일치']:,}",
                      f"{r['불일치']:,}",
                      f"{r.get('조인실패(경계밖)', 0):,}",
                      f"{r['일치율(%)']:.1f}%"],
                     center_cols=[1, 2, 3, 4, 5])

    # 불일치 상세
    has_detail = any(r.get("불일치상세") for r in geo_data["시설별결과"])
    if has_detail:
        doc.add_paragraph()
        _heading(doc, "4-1. 지오코딩 불일치 레코드 상세 (시설별 최대 20건)", level=2)
        for r in geo_data["시설별결과"]:
            if not r.get("불일치상세"):
                continue
            _para(doc, f"▶ {r['시설명']} ({r['불일치']}건)", bold=True, size=9)
            det_headers = ["시설명", "주소", "주소상 읍면동", "공간조인 읍면동"]
            det_table = doc.add_table(rows=1, cols=len(det_headers))
            det_table.style = "Table Grid"
            _header_row(det_table, det_headers, bg="5B7FC0")
            for rec in r["불일치상세"]:
                _add_row(det_table,
                         [rec.get("시설명", "-"),
                          rec.get("주소", "-"),
                          rec.get("주소상_읍면동", "-"),
                          rec.get("공간조인_읍면동", "-")],
                         center_cols=[2, 3])
    doc.add_paragraph()


def _section_summary(doc, count_data, dup_data, geo_data):
    _heading(doc, "5. 종합 의견", level=1)
    s = count_data["summary"]
    lines = [
        f"① [수량 비교] 전년 대비 총 {s['총증감']:+,}건 변동. "
        f"신규 시설 {len(s['신규시설'])}종, 추후 수령 후 분석 {len(s['추후수령후분석'])}종.",
        f"② [중복 제거] 총 {dup_data['총중복수']:,}건의 중복 레코드 탐지.",
        f"③ [지오코딩] 전체 일치율 {geo_data['전체일치율(%)']:.1f}% "
        f"({geo_data['전체일치']:,}/{geo_data['전체수량']:,}건 일치).",
    ]
    for line in lines:
        _para(doc, line, size=10)


# ── 메인 ─────────────────────────────────────────────────────

def generate_word_report(
    output_path: str,
    count_data: dict,
    dup_data: dict,
    geo_data: dict,
    prev_folder: str = "",
    curr_folder: str = "",
    emd_shp: str = "",
) -> str:
    doc = Document()

    # 제목
    title = doc.add_heading("시설 점(Point) 자료 검수 결과 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    _section_overview(doc, prev_folder, curr_folder, emd_shp)
    _section_count(doc, count_data)
    _section_duplicate(doc, dup_data)
    _section_geocode(doc, geo_data)
    _section_summary(doc, count_data, dup_data, geo_data)

    doc.save(output_path)
    return output_path
